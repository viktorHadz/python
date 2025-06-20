import re
import pandas as pd
import matplotlib.pyplot as plt
import textwrap
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Import configuration
from survey_config import (
    QUESTION_CATEGORIES, ACADEMIC_COLORS, DOC_STYLES, 
    EXCLUDE_COLUMNS, get_chart_color, should_use_horizontal_chart
)

def setup_document_styles(doc):
    # Setup professional document styling
    # Create custom styles
    styles = doc.styles
    
    # Main title style
    if 'CustomTitle' not in [s.name for s in styles]:
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = 18
        title_style.font.color.rgb = RGBColor(*DOC_STYLES['title_color'])
        title_style.font.bold = True
    
    # Section heading style
    if 'CustomHeading' not in [s.name for s in styles]:
        heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.size = 14
        heading_style.font.color.rgb = RGBColor(*DOC_STYLES['heading_color'])
        heading_style.font.bold = True

def create_executive_summary(doc, data, total_questions):
    # Create professional executive summary
    doc.add_heading("Executive Summary", level=1)
    
    summary_text = f"""
This comprehensive survey analysis examines student transition experiences at Teesside University London. 
The report analyzes {total_questions} questions across 9 key areas, based on responses from {len(data)} participants.

Key highlights include course selection motivations, academic preparedness, transition concerns, 
and demographic insights that inform student support strategies.
    """
    
    summary_para = doc.add_paragraph(summary_text.strip())
    summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def create_category_overview(doc):
    # Create table of contents style overview
    doc.add_heading("Survey Categories Overview", level=2)
    
    for i, (category_name, questions) in enumerate(QUESTION_CATEGORIES.items(), 1):
        overview_para = doc.add_paragraph()
        run = overview_para.add_run(f"{i}. {category_name}")
        run.bold = True
        run.font.color.rgb = RGBColor(*DOC_STYLES['heading_color'])
        overview_para.add_run(f" ({len(questions)} questions)")

def create_enhanced_chart(question, pivot_table, chart_path, question_index=0):
    # Create enhanced chart with professional styling
    # Determine chart orientation
    max_label_length = max(len(str(label)) for label in pivot_table[question])
    use_horizontal = should_use_horizontal_chart(question, len(pivot_table), max_label_length)
    
    # Special handling for questions 61 and 62 - always use horizontal with extra space
    question_num = question.split('.')[0]
    if question_num in ['61', '62']:
        use_horizontal = True
        fig_height = max(10, len(pivot_table) * 0.8)  # Extra height for these questions
        fig, ax = plt.subplots(figsize=(14, fig_height))  # Extra width too
    elif use_horizontal:
        fig_height = max(8, len(pivot_table) * 0.6)
        fig, ax = plt.subplots(figsize=(12, fig_height))
    else:
        fig, ax = plt.subplots(figsize=(12, 8))
    
    # Set style
    plt.style.use('default')
    fig.patch.set_facecolor('white')
    
    # Use consistent light blue color for all charts
    chart_color = '#87CEEB'  # Light blue from your academic palette
    text_color = '#2F2F2F'   # Almost black but softer
    
    if use_horizontal:
        # Horizontal bar chart
        bars = ax.barh(pivot_table[question], pivot_table["Count"], 
                      color=chart_color, alpha=0.8, edgecolor='white', linewidth=1)
        
        # Add data labels
        for bar in bars:
            width = bar.get_width()
            ax.text(width + 0.1, bar.get_y() + bar.get_height()/2.,
                    f'{int(width)}',
                    ha='left', va='center', fontsize=11, fontweight='bold', color=text_color)
        
        ax.set_xlabel("Count", fontsize=12, fontweight='bold', color=text_color)
        ax.set_ylabel("Response", fontsize=12, fontweight='bold', color=text_color)
        
        # Wrap labels for y-axis - extra space for questions 61 & 62
        if question_num in ['61', '62']:
            wrapped_labels = [textwrap.fill(str(label), 70) for label in pivot_table[question]]
            plt.subplots_adjust(left=0.45, right=0.95, top=0.9, bottom=0.1)
        else:
            wrapped_labels = [textwrap.fill(str(label), 50) for label in pivot_table[question]]
            plt.subplots_adjust(left=0.35, right=0.95, top=0.9, bottom=0.1)
        
        ax.set_yticks(range(len(wrapped_labels)))
        ax.set_yticklabels(wrapped_labels, fontsize=10, color=text_color)
        
    else:
        # Vertical bar chart
        bars = ax.bar(pivot_table[question], pivot_table["Count"], 
                     color=chart_color, alpha=0.8, edgecolor='white', linewidth=1)
        
        # Add data labels
        for bar in bars:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                    f'{int(height)}',
                    ha='center', va='bottom', fontsize=11, fontweight='bold', color=text_color)
        
        ax.set_xlabel("Response", fontsize=12, fontweight='bold', color=text_color)
        ax.set_ylabel("Count", fontsize=12, fontweight='bold', color=text_color)
        
        # Wrap labels
        wrapped_labels = [textwrap.fill(str(label), 15) for label in pivot_table[question]]
        ax.set_xticks(range(len(wrapped_labels)))
        ax.set_xticklabels(wrapped_labels, fontsize=10, color=text_color)
        
        plt.subplots_adjust(bottom=0.3, left=0.1, right=0.95, top=0.9)
    
    # Enhanced title - bigger and deep blue
    title_text = question.replace(question.split('.')[0] + '.', '').strip()
    ax.set_title(title_text, fontsize=16, fontweight='bold', 
                color=DOC_STYLES['title_color_hex'], pad=25)
    
    # Grid for better readability
    ax.grid(True, alpha=0.3, linestyle='--', color=text_color)
    ax.set_axisbelow(True)
    
    # Remove top and right spines for cleaner look
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color(text_color)
    ax.spines['bottom'].set_color(text_color)
    
    # Set tick colors
    ax.tick_params(colors=text_color)
    
    plt.savefig(chart_path, bbox_inches='tight', dpi=DOC_STYLES['chart_dpi'], 
                facecolor='white', edgecolor='none')
    plt.close()

def add_enhanced_analysis(doc, pivot_table, question):
    # Add enhanced analysis with insights
    total_responses = pivot_table['Count'].sum()
    top_response = pivot_table.iloc[0]
    
    # Calculate percentages
    pivot_table['Percentage'] = (pivot_table['Count'] / total_responses * 100).round(1)
    
    # Analysis paragraph
    analysis_para = doc.add_paragraph()
    analysis_para.add_run("Key Insights: ").bold = True
    
    insights = [
        f"Total responses: {total_responses}",
        f"Most common response: '{top_response[question]}' ({top_response['Percentage']}%)",
        f"Response diversity: {len(pivot_table)} different responses provided"
    ]
    
    if len(pivot_table) > 1:
        second_response = pivot_table.iloc[1]
        insights.append(f"Second most common: '{second_response[question]}' ({second_response['Percentage']}%)")
    
    analysis_para.add_run(" • ".join(insights))

def main():
    # Main function to generate the enhanced survey report
    # Load the dataset
    file_path = './transitions.csv'
    data = pd.read_csv(file_path)
    
    # Get actual questions (exclude metadata columns)
    all_questions = [col for col in data.columns if col not in EXCLUDE_COLUMNS]
    
    # Create enhanced document
    doc = Document()
    setup_document_styles(doc)
    
    # Main title
    title = doc.add_heading("Student Transition Experience Survey", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph("Teesside University London - Comprehensive Analysis Report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Executive summary
    create_executive_summary(doc, data, len(all_questions))
    doc.add_page_break()
    
    # Category overview
    create_category_overview(doc)
    doc.add_page_break()
    
    # Process each category
    question_counter = 0
    for category_name, category_questions in QUESTION_CATEGORIES.items():
        # Category header
        doc.add_heading(f"{category_name}", level=1)
        
        # Process questions in this category
        for question in category_questions:
            if question in data.columns:
                question_counter += 1
                
                # Count responses
                pivot_table = data[question].value_counts().reset_index()
                pivot_table.columns = [question, "Count"]
                
                # Question heading
                clean_question = question.replace(question.split('.')[0] + '.', '').strip()
                doc.add_heading(f"Q{question_counter}: {clean_question}", level=2)
                
                # Create data table
                table = doc.add_table(rows=min(len(pivot_table) + 1, 11), cols=3)  # Limit to top 10
                table.style = "Light Shading Accent 1"
                
                # Table headers
                table.cell(0, 0).text = "Response"
                table.cell(0, 1).text = "Count"
                table.cell(0, 2).text = "Percentage"
                
                # Calculate percentages
                total = pivot_table['Count'].sum()
                pivot_table['Percentage'] = (pivot_table['Count'] / total * 100).round(1)
                
                # Add data (top 10 only)
                for row_idx, row in pivot_table.head(10).iterrows():
                    table.cell(row_idx + 1, 0).text = str(row[question])[:50] + ("..." if len(str(row[question])) > 50 else "")
                    table.cell(row_idx + 1, 1).text = str(row["Count"])
                    table.cell(row_idx + 1, 2).text = f"{row['Percentage']}%"
                
                # Create chart
                sanitized_question = re.sub(r'[^\w\-_.]', '_', question)
                chart_path = f"{sanitized_question}_chart.png"
                
                create_enhanced_chart(question, pivot_table, chart_path, question_counter)
                
                # Add chart to document
                doc.add_picture(chart_path, width=Inches(DOC_STYLES['chart_width']))
                
                # Add enhanced analysis
                add_enhanced_analysis(doc, pivot_table, question)
                
                # Add space between questions
                doc.add_paragraph("")
        
        # Page break between categories
        if category_name != list(QUESTION_CATEGORIES.keys())[-1]:  # Not last category
            doc.add_page_break()
    
    # Save document
    output_path = "./Enhanced_Survey_Report.docx"
    doc.save(output_path)
    print(f"Enhanced survey report generated: {output_path}")
    print(f"Processed {question_counter} questions across {len(QUESTION_CATEGORIES)} categories")

if __name__ == "__main__":
    main()